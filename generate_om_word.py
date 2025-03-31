import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

def read_markdown_file(file_path):
    """Read markdown content from a file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
        # Skip the title line (# Section Name) as we'll add our own formatting
        if content.startswith('#'):
            content = '\n'.join(content.split('\n')[1:]).strip()
        return content

def format_section_title(document, title):
    """Add and format a section title"""
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue color
    
    # Add some space after the heading
    document.add_paragraph()

def add_bullet_list(document, items):
    """Add a bullet list to the document"""
    for item in items:
        p = document.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(item.strip())

def add_two_column_facts(document, facts):
    """Add facts in a two-column layout"""
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # Split facts into two columns
    facts_list = [fact.strip() for fact in facts.split('|')]
    mid_point = len(facts_list) // 2 + len(facts_list) % 2
    
    left_column = facts_list[:mid_point]
    right_column = facts_list[mid_point:]
    
    # Add rows to make both columns equal length
    max_rows = max(len(left_column), len(right_column))
    for i in range(max_rows):
        row = table.add_row()
        
        # Add left column item
        if i < len(left_column):
            row.cells[0].text = left_column[i]
            
        # Add right column item
        if i < len(right_column):
            row.cells[1].text = right_column[i]
    
    # Add some space after the table
    document.add_paragraph()

def main():
    # Create a new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('OFFER MEMORANDUM', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.bold = True
    
    # Add a page break after title page
    doc.add_page_break()
    
    # 1. Marketplace content with Key Valuation Points
    format_section_title(doc, "MARKETPLACE CONTENT")
    marketplace_content = read_markdown_file('output/marketplace_overview.md')
    doc.add_paragraph(marketplace_content)
    
    # Add Key Valuation Points as bullets at the end of marketplace content
    doc.add_heading('Key Valuation Points', level=2)
    facts_sheet = read_markdown_file('output/facts_sheet.md')
    facts_list = [fact.strip() for fact in facts_sheet.split('|')]
    add_bullet_list(doc, facts_list)
    
    # Add page break
    doc.add_page_break()
    
    # 2. Company Intro with KVP and Scaling Opportunities
    format_section_title(doc, "COMPANY INTRODUCTION")
    company_intro = read_markdown_file('output/company_intro.md')
    doc.add_paragraph(company_intro)
    
    # Add Key Valuation Points again
    doc.add_heading('Key Valuation Points', level=2)
    add_bullet_list(doc, facts_list)
    
    # Add Scaling Opportunities
    doc.add_heading('Scaling Opportunities', level=2)
    scaling_opps = read_markdown_file('output/scaling_opportunities.md')
    opps_list = [opp.strip().lstrip('> ') for opp in scaling_opps.split('\n') if opp.strip()]
    add_bullet_list(doc, opps_list)
    
    # Add page break
    doc.add_page_break()
    
    # 3. Company Overview
    format_section_title(doc, "COMPANY OVERVIEW")
    company_overview = read_markdown_file('output/company_overview.md')
    doc.add_paragraph(company_overview)
    
    # Add page break
    doc.add_page_break()
    
    # 4. Facts Sheet as two columns
    format_section_title(doc, "FACTS SHEET")
    add_two_column_facts(doc, facts_sheet)
    
    # Add page break
    doc.add_page_break()
    
    # 5. About Us
    format_section_title(doc, "ABOUT US")
    about_us = read_markdown_file('output/about_us.md')
    doc.add_paragraph(about_us)
    
    # Add page break
    doc.add_page_break()
    
    # 6. Key Methods to Scale
    format_section_title(doc, "KEY METHODS TO SCALE")
    scaling_strategy = read_markdown_file('output/scaling_strategy.md')
    doc.add_paragraph(scaling_strategy)
    
    # Add page break
    doc.add_page_break()
    
    # 7. Industry Overview
    format_section_title(doc, "INDUSTRY OVERVIEW")
    industry_overview = read_markdown_file('output/industry_overview.md')
    doc.add_paragraph(industry_overview)
    
    # Save the document
    output_file = 'Final_Offer_Memorandum.docx'
    doc.save(output_file)
    print(f"Offer Memorandum successfully generated as '{output_file}'")

if __name__ == "__main__":
    main() 