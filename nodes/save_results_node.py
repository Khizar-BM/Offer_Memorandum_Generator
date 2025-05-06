from state import GraphState
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def format_section_title(document, title):
    """Add and format a section title"""
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue color
    
    # Add some space after the heading
    document.add_paragraph()


def add_bullet_list(document, items):
    """Add a bullet list to the document"""
    for item in items:
        p = document.add_paragraph()
        p.style = 'List Bullet'
        
        # Handle bold text formatting in list items
        item_text = item.strip()
        if '**' in item_text:
            # Split by ** markers and handle bold formatting
            parts = re.split(r'(\*\*)', item_text)
            is_bold = False
            for part in parts:
                if part == '**':
                    is_bold = not is_bold
                elif part:
                    run = p.add_run(part)
                    run.bold = is_bold
        else:
            p.add_run(item_text)


def add_two_column_facts(document, facts):
    """Add facts in a two-column layout"""
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # Split facts into two columns
    facts_list = [fact.strip() for fact in facts.split('|')]
    mid_point = len(facts_list) // 2 + len(facts_list) % 2
    
    left_column = facts_list[:mid_point]
    right_column = facts_list[mid_point:]
    
    # Add rows to make both columns equal length
    max_rows = max(len(left_column), len(right_column))
    for i in range(max_rows):
        row = table.add_row()
        
        # Add left column item with bold formatting
        if i < len(left_column):
            cell = row.cells[0]
            fact_text = left_column[i]
            if '**' in fact_text:
                parts = re.split(r'(\*\*)', fact_text)
                paragraph = cell.paragraphs[0]
                is_bold = False
                for part in parts:
                    if part == '**':
                        is_bold = not is_bold
                    elif part:
                        run = paragraph.add_run(part)
                        run.bold = is_bold
            else:
                cell.text = fact_text
            
        # Add right column item with bold formatting
        if i < len(right_column):
            cell = row.cells[1]
            fact_text = right_column[i]
            if '**' in fact_text:
                parts = re.split(r'(\*\*)', fact_text)
                paragraph = cell.paragraphs[0]
                is_bold = False
                for part in parts:
                    if part == '**':
                        is_bold = not is_bold
                    elif part:
                        run = paragraph.add_run(part)
                        run.bold = is_bold
            else:
                cell.text = fact_text
    
    # Add some space after the table
    document.add_paragraph()


def format_rich_text(document, content):
    """Format text with proper headings and paragraphs"""
    paragraphs = content.split('\n')
    i = 0
    
    while i < len(paragraphs):
        paragraph = paragraphs[i].strip()
        
        # Skip empty paragraphs
        if not paragraph:
            i += 1
            continue
        
        # Check for headings and handle $1 placeholders that should be headings
        if paragraph == '$1':
            # This is a placeholder for a heading - check next paragraph to use as heading
            if i + 1 < len(paragraphs) and paragraphs[i+1].strip():
                heading_text = paragraphs[i+1].strip()
                heading = document.add_heading(heading_text, level=2)
                for run in heading.runs:
                    run.font.bold = True
                # Skip the next paragraph since we used it as heading
                i += 2
                continue
            else:
                # Just skip this placeholder if there's no text after it
                i += 1
                continue
        elif paragraph.startswith('# '):
            # Level 1 heading
            heading_text = paragraph[2:].strip()
            heading = document.add_heading(heading_text, level=1)
            for run in heading.runs:
                run.font.bold = True
        elif paragraph.startswith('## '):
            # Level 2 heading
            heading_text = paragraph[3:].strip()
            heading = document.add_heading(heading_text, level=2)
            for run in heading.runs:
                run.font.bold = True
        elif paragraph.startswith('### '):
            # Level 3 heading
            heading_text = paragraph[4:].strip()
            heading = document.add_heading(heading_text, level=2)
            for run in heading.runs:
                run.font.bold = True
        elif paragraph.startswith('#### '):
            # Level 4 heading
            heading_text = paragraph[5:].strip()
            heading = document.add_heading(heading_text, level=3)
            for run in heading.runs:
                run.font.bold = True
        elif paragraph.startswith('**Business Broker Takeaways:**'):
            # Special formatting for Business Broker Takeaways
            takeaway_heading = document.add_heading('Business Broker Takeaways:', level=2)
            for run in takeaway_heading.runs:
                run.font.bold = True
                
            # Process numbered points that follow
            j = i + 1
            while j < len(paragraphs) and re.match(r'^\d+\.\s+\*\*', paragraphs[j].strip()):
                point = paragraphs[j].strip()
                # Extract the number and title
                match = re.match(r'(\d+)\.\s+\*\*([^:]+):\*\*\s*(.*)', point)
                if match:
                    number, title, content = match.groups()
                    p = document.add_paragraph()
                    p.style = 'List Number'
                    run = p.add_run(f"{title}: ")
                    run.bold = True
                    p.add_run(content)
                else:
                    # If the pattern doesn't match exactly, just add as is
                    p = document.add_paragraph()
                    p.style = 'List Number'
                    p.add_run(re.sub(r'^\d+\.\s+', '', point))
                j += 1
            
            # Skip the processed points
            i = j - 1
        else:
            # Handle bold text formatting (**text**)
            if '**' in paragraph:
                p = document.add_paragraph()
                # Split by ** markers
                parts = re.split(r'(\*\*)', paragraph)
                
                is_bold = False
                for part in parts:
                    if part == '**':
                        is_bold = not is_bold
                    elif part:
                        run = p.add_run(part)
                        run.bold = is_bold
            else:
                # Regular paragraph
                document.add_paragraph(paragraph)
        
        i += 1


def generate_word_document(om_sections, output_dir, selected_broker="Website Closers", main_company_name="Main Company"):
    """
    Generate a Word document from the OM sections.
    Also handles replacing 'Website Closers' with 'SellerForce' if needed.
    """
    # Create a new Word document
    doc = Document()
    
    # Preprocess all sections to fix common formatting issues
    for section_key in om_sections:
        content = om_sections[section_key]
        if isinstance(content, str):
            # Replace $1 markers followed by text with proper headings
            content = re.sub(r'\$1\s*\n([^\n]+)', r'## \1', content)
            
            # Ensure proper Markdown formatting for bold text
            content = re.sub(r'\*\*([^*]+)\*\*\s*(?!\n)', r'**\1**\n', content)
            
            # Clean up any inconsistent heading formatting
            content = re.sub(r'##\s+([^\n]+)', r'## \1', content)
            
            # Update the section with cleaned content
            om_sections[section_key] = content
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('OFFER MEMORANDUM', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.bold = True
    
    # Add broker name subtitle
    broker_subtitle = doc.add_paragraph()
    broker_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    broker_run = broker_subtitle.add_run(selected_broker)
    broker_run.font.size = Pt(16)
    broker_run.font.bold = True
    
    # Add main company name
    company_subtitle = doc.add_paragraph()
    company_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_subtitle.add_run(main_company_name)
    company_run.font.size = Pt(20)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
    
    # Add a page break after title page
    doc.add_page_break()
    
    # 1. Marketplace content with Key Valuation Points
    format_section_title(doc, "MARKETPLACE CONTENT")
    format_rich_text(doc, om_sections.get('marketplace_overview', ''))
    
    # Add Key Valuation Points as bullets at the end of marketplace content
    doc.add_heading('Key Valuation Points', level=2)
    facts_sheet = om_sections.get('facts_sheet', '')
    facts_list = [fact.strip() for fact in facts_sheet.split('|')]
    add_bullet_list(doc, facts_list)
    
    # Add page break
    doc.add_page_break()
    
    # 2. Company Intro with KVP and Scaling Opportunities
    format_section_title(doc, "COMPANY INTRODUCTION")
    format_rich_text(doc, om_sections.get('company_intro', ''))
    
    # Add Key Valuation Points again
    doc.add_heading('Key Valuation Points', level=2)
    add_bullet_list(doc, facts_list)
    
    # Add Scaling Opportunities
    doc.add_heading('Scaling Opportunities', level=2)
    scaling_opps = om_sections.get('scaling_opportunities', '')
    # Parse scaling opportunities, handling both bullet points and regular text
    opps_list = []
    for opp in scaling_opps.split('\n'):
        opp = opp.strip()
        if opp:
            # Remove markdown bullet points and '>' quote markers
            clean_opp = re.sub(r'^[>*\-•]+\s*', '', opp)
            opps_list.append(clean_opp)
    add_bullet_list(doc, opps_list)
    
    # Add page break
    doc.add_page_break()
    
    # 3. Company Overview
    format_section_title(doc, "COMPANY OVERVIEW")
    format_rich_text(doc, om_sections.get('company_overview', ''))
    
    # 4. Company Summary after Company Overview
    format_section_title(doc, "COMPANY SUMMARY")
    format_rich_text(doc, om_sections.get('company_summary', ''))
    
    # Add page break
    doc.add_page_break()
    
    # 5. Facts Sheet as two columns
    format_section_title(doc, "FACTS SHEET")
    add_two_column_facts(doc, facts_sheet)
    
    # Add page break
    doc.add_page_break()
    
    # 6. About Us
    format_section_title(doc, "ABOUT US")
    format_rich_text(doc, om_sections.get('about_us', ''))
    
    # Add page break
    doc.add_page_break()
    
    # 7. Key Methods to Scale
    format_section_title(doc, "KEY METHODS TO SCALE")
    # Use the preprocessed scaling strategy content
    scaling_strategy = om_sections.get('scaling_strategy', '')
    format_rich_text(doc, scaling_strategy)
    
    # Add page break
    doc.add_page_break()
    
    # 8. Industry Overview
    format_section_title(doc, "INDUSTRY OVERVIEW")
    format_rich_text(doc, om_sections.get('industry_overview', ''))
    
    # Save the document
    # Use a fixed filename for the saved file
    output_file = f"{output_dir}/Offer_Memorandum.docx"
    
    doc.save(output_file)
    print(f"Word document generated successfully: {output_file}")
    
    return output_file


def save_results_node(state: GraphState) -> GraphState:
    """Save the generated OM sections to files and create a Word document"""
    om_sections = state.get("om_sections", {})
    selected_broker = state.get("selected_broker", "Website Closers")
    main_company_name = state.get("main_company_name", "Main Company")
    print(f"Generating OM Document for {selected_broker}, Company: {main_company_name}...")
    
    if not om_sections:
        return {**state, "error": "No OM sections to save"}
    
    try:
        # Replace "Website Closers" with "SellerForce" if SellerForce is selected
        if selected_broker == "SellerForce":
            print("Replacing 'Website Closers' with 'SellerForce' in all content...")
            # Create a copy of the sections dict with modified content
            modified_sections = {}
            for section_name, content in om_sections.items():
                # Replace all instances of "Website Closers" with "SellerForce"
                modified_content = content.replace("Website Closers", "SellerForce")
                modified_sections[section_name] = modified_content
            
            # Update the sections for further processing
            om_sections = modified_sections
        
        # Create output directory
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
        
       
             
        # Generate Word document
        word_doc_path = generate_word_document(om_sections, output_dir, selected_broker, main_company_name)
        
        print(f"Offer Memorandum generated successfully in the '{output_dir}' directory.")
        return {**state, "word_document_path": word_doc_path, "error": None}
    except Exception as e:
        return {**state, "error": f"Failed to save results: {str(e)}"}

